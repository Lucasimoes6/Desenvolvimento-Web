from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Configuração de página (A4, margens ABNT) ──────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2)
section.top_margin    = Cm(3)
section.bottom_margin = Cm(2)

# ── Estilos base ──────────────────────────────────────────────────────────
FONTE = "Times New Roman"

def set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = FONTE
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, size=14, bold=True, upper=False, center=False, page_break_before=False):
    p = doc.add_paragraph()
    if page_break_before:
        p.runs  # ensure paragraph exists
        from docx.oxml import OxmlElement
        pPr = p._p.get_or_add_pPr()
        pb = OxmlElement('w:pageBreakBefore')
        pb.set(qn('w:val'), '1')
        pPr.append(pb)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper() if upper else text)
    set_font(run, size=size, bold=bold)
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after  = Pt(6)
    return p

def add_body(doc, text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    if indent:
        pf.first_line_indent = Cm(1.5)
    pf.space_after = Pt(6)
    pf.line_spacing = Pt(18)
    run = p.add_run(text)
    set_font(run, size=12)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(2)
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text)
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p

def add_code_block(doc, code):
    for line in code.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
        # fundo cinza claro via shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F4F4F4')
        pPr.append(shd)

def add_figure_placeholder(doc, label):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after  = Pt(4)
    pf.border = None
    run = p.add_run(f"[ {label} ]")
    run.font.name = FONTE
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

def add_figure_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_after = Pt(14)
    run = p.add_run(text)
    run.font.name = FONTE
    run.font.size = Pt(10)
    run.font.italic = True

def page_break(doc):
    doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# CAPA
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(30)
r = p.add_run("Facens")
set_font(r, size=28, bold=True, color=(0x1A, 0x56, 0xA0))

# Tabela de integrantes
table = doc.add_table(rows=6, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["N°", "NOME", "E-mail", "Telefone"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

integrantes = [
    ("247926", "Lucas Machado Simoes"),
    ("249313", "Murilo Guazzelli da Silveira"),
    ("247771", "Guilherme Ferrari Assad Crudo"),
    ("247963", "Fabricio Guima Giacomelli"),
    ("248228", "Murilo de Oliveira Souza"),
]
for i, (num, nome) in enumerate(integrantes, start=1):
    row = table.rows[i]
    row.cells[0].text = num
    row.cells[1].text = nome
    for c in row.cells:
        for para in c.paragraphs:
            for run in para.runs:
                run.font.name = FONTE
                run.font.size = Pt(11)
                run.font.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ATIVIDADE PRÁTICA 9:")
set_font(r, size=14, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Desenvolvimento de API REST com Express.js e Front-End em React")
set_font(r, size=12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("(v1.0.0)")
set_font(r, size=11)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SOROCABA, SP\n2026")
set_font(r, size=12)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# SUMÁRIO
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "SUMÁRIO", size=14, center=True, upper=False)

sumario = [
    ("1 PROJETO BACK-END (EXPRESS.JS)", "3"),
    ("2 PROJETO FRONT-END (REACT)", "4"),
    ("3 DEPLOY – RENDER E VERCEL", "5"),
    ("4 DOCUMENTAÇÃO DA API NO POSTMAN", "6"),
    ("REFERÊNCIAS", "7"),
]
for titulo, pg in sumario:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    tab_stops = p.paragraph_format.tab_stops
    run = p.add_run(f"{titulo}\t{pg}")
    set_font(run, size=12)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# SEÇÃO 1 – BACK-END
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1 PROJETO BACK-END (EXPRESS.JS)", size=12, upper=False)

add_body(doc, (
    "O primeiro projeto desenvolvido consiste em uma API REST para gerenciamento de notas, construída com o "
    "framework Express.js sobre a plataforma Node.js. A aplicação expõe cinco endpoints que cobrem todas as "
    "operações CRUD (Create, Read, Update, Delete), permitindo criar, listar, buscar por ID, editar e excluir notas. "
    "Os dados são persistidos em um arquivo data.json, lido e gravado de forma síncrona via módulo nativo fs do "
    "Node.js. O middleware body-parser é utilizado para interpretar o corpo das requisições no formato JSON, "
    "enquanto um middleware de CORS personalizado libera o acesso externo de qualquer origem."
))

add_body(doc, (
    "A estrutura do projeto é composta por dois arquivos principais: server.js, responsável por toda a lógica "
    "do servidor, e data.json, que armazena as notas. O servidor escuta na porta 3000 por padrão, com suporte "
    "à variável de ambiente PORT para facilitar o deploy em nuvem. O código foi versionado com Git e hospedado "
    "no GitHub, e o deploy foi realizado na plataforma Render, tornando a API acessível publicamente via HTTPS."
))

add_bullet(doc, " [inserir link do repositório backend]", bold_prefix="GitHub: ")
add_bullet(doc, " [inserir URL do Render, ex: https://seu-projeto.onrender.com]", bold_prefix="Render: ")

p = doc.add_paragraph()
r = p.add_run("A seguir, o código completo do arquivo server.js:")
set_font(r, size=12)
p.paragraph_format.space_after = Pt(4)

server_code = """// Importa o Express, Body-Parser e FS
const express = require('express');
const bodyParser = require('body-parser');
const fs = require('fs');

const app = express();
const PORT = process.env.PORT || 3000;
const FILE = 'data.json';

app.use(bodyParser.json());

app.use((req, res, next) => {
  res.header('Access-Control-Allow-Origin', '*');
  res.header('Access-Control-Allow-Methods', 'GET, POST, PUT, DELETE, OPTIONS');
  res.header('Access-Control-Allow-Headers', 'Content-Type');
  if (req.method === 'OPTIONS') return res.sendStatus(204);
  next();
});

function readNotes() {
  try { return JSON.parse(fs.readFileSync(FILE)); }
  catch { return []; }
}

function saveNotes(notes) {
  fs.writeFileSync(FILE, JSON.stringify(notes, null, 2));
}

// GET - Listar todas as notas
app.get('/api/notes', (req, res) => {
  res.json(readNotes());
});

// GET - Nota por ID
app.get('/api/notes/:id', (req, res) => {
  const note = readNotes().find(n => n.id === req.params.id);
  if (note) res.json(note);
  else res.status(404).json({ erro: 'Nota não encontrada' });
});

// POST - Criar nota
app.post('/api/notes', (req, res) => {
  const { titulo, texto } = req.body;
  if (!titulo || !texto)
    return res.status(400).json({ erro: 'Os campos titulo e texto são obrigatórios' });
  const notes = readNotes();
  const novaNota = { id: Date.now().toString(), titulo, texto, criadoEm: new Date().toISOString() };
  notes.push(novaNota);
  saveNotes(notes);
  res.status(201).json(novaNota);
});

// PUT - Editar nota
app.put('/api/notes/:id', (req, res) => {
  const { titulo, texto } = req.body;
  if (!titulo || !texto)
    return res.status(400).json({ erro: 'Os campos titulo e texto são obrigatórios' });
  const notes = readNotes();
  const index = notes.findIndex(n => n.id === req.params.id);
  if (index >= 0) {
    notes[index] = { ...notes[index], titulo, texto, atualizadoEm: new Date().toISOString() };
    saveNotes(notes);
    res.json(notes[index]);
  } else {
    res.status(404).json({ erro: 'Nota não encontrada' });
  }
});

// DELETE - Excluir nota
app.delete('/api/notes/:id', (req, res) => {
  const notes = readNotes();
  if (!notes.find(n => n.id === req.params.id))
    return res.status(404).json({ erro: 'Nota não encontrada' });
  saveNotes(notes.filter(n => n.id !== req.params.id));
  res.status(204).send();
});

app.listen(PORT, () => console.log(`Servidor rodando em http://localhost:${PORT}`));"""

add_code_block(doc, server_code)

add_figure_placeholder(doc, "Print do terminal com o servidor rodando / Print da API respondendo no navegador")
add_figure_caption(doc, "Figura 1 – API REST rodando localmente na porta 3000.\nFonte: Elaborado pelo autor (2026).")

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# SEÇÃO 2 – FRONT-END
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2 PROJETO FRONT-END (REACT)", size=12, upper=False)

add_body(doc, (
    "O segundo projeto consiste em uma interface web desenvolvida com a biblioteca React, inicializada com a "
    "ferramenta Vite. A aplicação consome todos os endpoints da API REST criada no back-end, permitindo ao usuário "
    "realizar as quatro operações CRUD diretamente pelo navegador: criar novas notas preenchendo título e texto, "
    "visualizar todas as notas salvas em cartões individuais, editar o conteúdo de uma nota existente e excluí-la "
    "após confirmação."
))

add_body(doc, (
    "Toda a lógica foi implementada em um único componente funcional App, utilizando os hooks useState e useEffect "
    "do React. O hook useEffect dispara a busca inicial das notas assim que o componente é montado, enquanto o "
    "estado local gerencia os campos do formulário, a lista de notas, o identificador da nota em edição e as "
    "mensagens de feedback ao usuário. A estilização foi feita via objetos JavaScript inline, sem dependência de "
    "bibliotecas externas de UI. O código foi versionado com Git e hospedado no GitHub, com deploy realizado na "
    "plataforma Vercel."
))

add_bullet(doc, " [inserir link do repositório frontend]", bold_prefix="GitHub: ")
add_bullet(doc, " [inserir URL da Vercel, ex: https://projeto-notas.vercel.app]", bold_prefix="Vercel: ")

p = doc.add_paragraph()
r = p.add_run("A seguir, o código principal do componente App.jsx:")
set_font(r, size=12)
p.paragraph_format.space_after = Pt(4)

app_code = """import { useState, useEffect } from "react";

const API_URL = "https://SEU-PROJETO.onrender.com/api/notes";

export default function App() {
  const [notas, setNotas] = useState([]);
  const [titulo, setTitulo] = useState("");
  const [texto, setTexto] = useState("");
  const [editandoId, setEditandoId] = useState(null);
  const [mensagem, setMensagem] = useState("");
  const [carregando, setCarregando] = useState(false);

  useEffect(() => { buscarNotas(); }, []);

  async function buscarNotas() {
    setCarregando(true);
    try {
      const res = await fetch(API_URL);
      setNotas(await res.json());
    } catch { exibirMensagem("Erro ao carregar notas.", "erro"); }
    finally { setCarregando(false); }
  }

  function exibirMensagem(msg, tipo = "sucesso") {
    setMensagem({ texto: msg, tipo });
    setTimeout(() => setMensagem(""), 3000);
  }

  async function salvarNota() {
    if (!titulo.trim() || !texto.trim())
      return exibirMensagem("Preencha título e texto.", "erro");
    try {
      if (editandoId) {
        await fetch(`${API_URL}/${editandoId}`, {
          method: "PUT",
          headers: { "Content-Type": "application/json" },
          body: JSON.stringify({ titulo, texto }),
        });
        exibirMensagem("Nota atualizada com sucesso!");
      } else {
        await fetch(API_URL, {
          method: "POST",
          headers: { "Content-Type": "application/json" },
          body: JSON.stringify({ titulo, texto }),
        });
        exibirMensagem("Nota criada com sucesso!");
      }
      setTitulo(""); setTexto(""); setEditandoId(null);
      buscarNotas();
    } catch { exibirMensagem("Erro ao salvar nota.", "erro"); }
  }

  async function excluirNota(id) {
    if (!window.confirm("Deseja excluir esta nota?")) return;
    try {
      await fetch(`${API_URL}/${id}`, { method: "DELETE" });
      exibirMensagem("Nota excluída com sucesso!");
      buscarNotas();
    } catch { exibirMensagem("Erro ao excluir nota.", "erro"); }
  }

  // ... retorno JSX com formulário e listagem de cartões de nota
}"""

add_code_block(doc, app_code)

add_figure_placeholder(doc, "Print da aplicação React no navegador com notas listadas")
add_figure_caption(doc, "Figura 2 – Interface React do Gerenciador de Notas.\nFonte: Elaborado pelo autor (2026).")

add_figure_placeholder(doc, "Print mostrando criação de nota com mensagem de sucesso")
add_figure_caption(doc, "Figura 3 – Criação de nota com feedback de sucesso.\nFonte: Elaborado pelo autor (2026).")

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# SEÇÃO 3 – DEPLOY
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3 DEPLOY – RENDER E VERCEL", size=12, upper=False)

add_heading(doc, "3.1 Back-End no Render", size=12, bold=True, upper=False)

add_body(doc, (
    "O deploy do back-end foi realizado na plataforma Render, que oferece hospedagem gratuita para serviços "
    "Node.js com integração direta a repositórios GitHub. Para efetuar o deploy, o repositório do back-end foi "
    "publicado no GitHub e conectado ao Render. Na configuração do serviço, o comando de build foi definido como "
    "npm install e o comando de inicialização como node server.js. Após o deploy, a API passou a estar disponível "
    "em uma URL pública no formato https://[nome-do-projeto].onrender.com, possibilitando que o front-end a consuma "
    "remotamente, sem necessidade de executar o servidor local."
))

add_figure_placeholder(doc, "Print do painel do Render com o serviço em execução (status: Live)")
add_figure_caption(doc, "Figura 4 – Deploy do back-end na plataforma Render.\nFonte: Elaborado pelo autor (2026).")

add_heading(doc, "3.2 Front-End na Vercel", size=12, bold=True, upper=False)

add_body(doc, (
    "O deploy do front-end foi realizado na plataforma Vercel, especializada em aplicações JavaScript modernas. "
    "O repositório do projeto React foi publicado no GitHub e importado diretamente na Vercel, que identificou "
    "automaticamente o framework Vite e configurou os comandos de build (npm run build) e o diretório de saída "
    "(dist). Antes do deploy, a constante API_URL no arquivo App.jsx foi atualizada com a URL pública gerada pelo "
    "Render, garantindo a comunicação correta entre front-end e back-end em produção."
))

add_bullet(doc, " [inserir URL do Render]", bold_prefix="Back-End (Render): ")
add_bullet(doc, " [inserir URL da Vercel]", bold_prefix="Front-End (Vercel): ")

add_figure_placeholder(doc, "Print do painel da Vercel com o projeto implantado e URL gerada")
add_figure_caption(doc, "Figura 5 – Deploy do front-end na plataforma Vercel.\nFonte: Elaborado pelo autor (2026).")

add_figure_placeholder(doc, "Print da aplicação acessada pela URL pública da Vercel, com notas carregadas da API no Render")
add_figure_caption(doc, "Figura 6 – Aplicação em produção integrando Render e Vercel.\nFonte: Elaborado pelo autor (2026).")

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# SEÇÃO 4 – POSTMAN
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4 DOCUMENTAÇÃO DA API NO POSTMAN", size=12, upper=False)

add_body(doc, (
    "A documentação da API foi elaborada no Postman, ferramenta amplamente utilizada para testar e documentar "
    "requisições HTTP. Foi criada uma coleção denominada API Notas contendo cinco requisições, uma para cada "
    "endpoint disponível. Cada requisição foi configurada com o método HTTP adequado, a URL correta, os cabeçalhos "
    "necessários e, quando aplicável, o corpo da requisição no formato JSON. Além disso, foram adicionadas "
    "descrições explicando a finalidade de cada endpoint e os códigos de resposta esperados."
))

add_body(doc, "As operações documentadas na coleção são:", indent=False)

endpoints = [
    ("GET /api/notes", "Lista todas as notas. Retorna status 200 com array JSON."),
    ("GET /api/notes/:id", "Retorna uma nota específica pelo ID. Status 200 se encontrada, 404 caso contrário."),
    ("POST /api/notes", "Cria uma nova nota. Requer titulo e texto no corpo. Retorna status 201 com a nota criada."),
    ("PUT /api/notes/:id", "Atualiza uma nota existente. Requer titulo e texto. Status 200 ou 404."),
    ("DELETE /api/notes/:id", "Exclui uma nota pelo ID. Retorna status 204 sem corpo, ou 404 se não encontrada."),
]
for ep, desc in endpoints:
    add_bullet(doc, f" – {desc}", bold_prefix=ep)

add_bullet(doc, " [inserir link público da coleção no Postman]", bold_prefix="Link da coleção Postman: ")

add_figure_placeholder(doc, "Print do Postman – GET /api/notes com resposta 200 e lista de notas")
add_figure_caption(doc, "Figura 7 – Requisição GET para listagem de notas no Postman.\nFonte: Elaborado pelo autor (2026).")

add_figure_placeholder(doc, "Print do Postman – POST /api/notes com body JSON e resposta 201")
add_figure_caption(doc, "Figura 8 – Requisição POST para criação de nota no Postman.\nFonte: Elaborado pelo autor (2026).")

add_figure_placeholder(doc, "Print do Postman – PUT /api/notes/:id com body JSON e resposta 200")
add_figure_caption(doc, "Figura 9 – Requisição PUT para edição de nota no Postman.\nFonte: Elaborado pelo autor (2026).")

add_figure_placeholder(doc, "Print do Postman – DELETE /api/notes/:id com resposta 204")
add_figure_caption(doc, "Figura 10 – Requisição DELETE para exclusão de nota no Postman.\nFonte: Elaborado pelo autor (2026).")

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# REFERÊNCIAS
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "REFERÊNCIAS", size=12, upper=False)

refs = [
    ("OLIVEIRA, Cláudio Luís Vieira. ", "Node.js: programe de forma rápida e prática.", " São Paulo: Expressa, 2021. ISBN 9786558110217."),
    ("FERREIRA, Arthur Gonçalves. ", "Interface de programação de aplicações (API) e web services.", " São Paulo: Platos Soluções Educacionais, 2021. ISBN 9786553560338."),
    ("MILETTO, Evandro Manara; DE CASTRO BERTAGNOLLI, Silvia. ", "Desenvolvimento de Software II.", " Introdução ao Desenvolvimento Web com HTML, CSS, JavaScript e PHP. Bookman, 2014."),
    ("RANJAN, Alok; SINHA, Abhilasha; BATTEWAD, Ranjit. ", "JavaScript for Modern Web Development.", " BPB Publications, 2020."),
    ("EXPRESS.JS. ", "Express – Fast, unopinionated, minimalist web framework for Node.js.", " Disponível em: https://expressjs.com. Acesso em: maio 2026."),
    ("RENDER. ", "Cloud Application Hosting for Developers.", " Disponível em: https://render.com. Acesso em: maio 2026."),
    ("VERCEL. ", "Develop. Preview. Ship.", " Disponível em: https://vercel.com. Acesso em: maio 2026."),
    ("POSTMAN. ", "The Collaboration Platform for API Development.", " Disponível em: https://www.postman.com. Acesso em: maio 2026."),
]

for before, bold_part, after in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent = Cm(1.5)
    pf.first_line_indent = Cm(-1.5)
    pf.space_after = Pt(8)
    pf.line_spacing = Pt(18)
    r1 = p.add_run(before)
    set_font(r1)
    r2 = p.add_run(bold_part)
    set_font(r2, bold=True)
    r3 = p.add_run(after)
    set_font(r3)

# Salva
output = r"c:\Users\lucas\OneDrive\Facens\5 Semestre\Desenvolvimento Web\aula-09-backend\Relatorio-Atividade9.docx"
doc.save(output)
print(f"Arquivo salvo em: {output}")
